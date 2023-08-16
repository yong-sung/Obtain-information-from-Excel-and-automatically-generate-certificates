from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

# 엑셀 파일을 불러와 워크북과 시트를 설정
load_wb = load_workbook(r"12. 엑셀의 정보를 불러와 수료증 자동생성_X\수료증명단.xlsx")

# 해당 워크시트에 접근
load_ws =load_wb.active

# 이름, 생년월일, 호 리스트를 초기화
name_list = []
birthday_list = []
ho_list = []

# 엑셀 파일로부터 데이터를 추출해 리스트에 저장
for i in range(1,load_ws.max_row + 1):
    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)
    
print(name_list)
print(birthday_list)
print(ho_list)

# 수료증 자동 생성 작업을 위한 루프 시작
for i in range(len(name_list)):
    doc = docx.Document(r'12. 엑셀의 정보를 불러와 수료증 자동생성_X\수료증양식.docx') # 수료증 양식 문서를 불러옴
    
    # 수료증 양식 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)
    
    # 수료증 내용을 추가
    para = doc.add_paragraph() # 문서에 새로운 문단을 추가하는 메서드
    run = para.add_run('\n\n') # 새로 추가한 문단에 텍스트 런을 추가
    
    # 처리할 로직 추가: ho_list가 None인 경우 빈 문자열로 처리
    ho_value = ho_list[i] if ho_list[i] is not None else ""
    
    # 호 정보를 가져와서 유효성 검사하고 추가
    ho_value = ho_list[i] if ho_list[i] and ho_list[i].startswith('20') else ""
    if not ho_value:
        print(f"Invalid ho_value for {name_list[i]}: {ho_list[i]}")
        continue  # 유효하지 않은 값은 처리하지 않고 건너뜁니다.
    
    run = para.add_run('              제 '+ ho_list[i] +' 호\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    # 수료증 제목을 추가
    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수  료  증') 
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 수료증 내용을 추가
    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        성       명: ' + name_list[i] +'\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        생 년 월 일: ' + birthday_list[i] +'\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20) 
    run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('2021.09.19') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('파이썬교육기관장') 
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 문서를 저장하고 PDF로 변환
    doc.save('12. 엑셀의 정보를 불러와 수료증 자동생성_X\\'+name_list[i]+'.docx')
    convert('12. 엑셀의 정보를 불러와 수료증 자동생성_X\\'+name_list[i]+'.docx',
            '12. 엑셀의 정보를 불러와 수료증 자동생성_X\\'+name_list[i]+'.pdf')